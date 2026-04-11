import docx
import pandas as pd

# 读取Word文档
doc = docx.Document('c:\\Users\\22307\\Desktop\\ai-medicine-assistant-main\\第二模块大纲.docx')

print(f"文档中包含 {len(doc.tables)} 个表格")

# 打印每个表格的内容
for i, table in enumerate(doc.tables):
    print(f"\n表格 {i+1}:")
    print("=" * 50)
    
    # 将表格转换为DataFrame
    data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        data.append(row_data)
    
    df = pd.DataFrame(data)
    print(df)
    print("\n" + "-" * 50)