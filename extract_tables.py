from docx import Document
import json

def extract_tables_and_text(file_path):
    """
    提取Word文档中的表格和文本
    """
    doc = Document(file_path)

    # 提取表格
    tables = []
    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        tables.append({
            "index": table_idx,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "data": table_data
        })

    # 提取段落
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    # 保存结果
    result = {
        "total_paragraphs": len(paragraphs),
        "total_tables": len(tables),
        "paragraphs": paragraphs,
        "tables": tables
    }

    with open('docx_full_content.json', 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"提取完成：{len(paragraphs)}个段落，{len(tables)}个表格")
    print("已保存到 docx_full_content.json")

    # 显示表格信息
    print("\n=== 表格信息 ===")
    for i, table in enumerate(tables[:5]):  # 只显示前5个表格
        print(f"表格 {i+1}: {table['rows']}行 x {table['cols']}列")
        print(f"第一行内容: {table['data'][0]}")

    return result

if __name__ == "__main__":
    file_path = "2025年国家执业药师职业资格考试大纲.docx"
    extract_tables_and_text(file_path)