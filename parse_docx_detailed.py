from docx import Document
import json

def parse_docx_content():
    """解析DOCX文件中的表格内容"""
    doc = Document('2025年国家执业药师职业资格考试大纲.docx')
    
    print(f"=== 解析DOCX文件 ===")
    print(f"总表格数: {len(doc.tables)}")
    
    # 解析所有表格
    all_tables_data = []
    for table_idx, table in enumerate(doc.tables):
        print(f"\n表格 {table_idx}:")
        table_data = []
        
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data.append(row_data)
            
            if row_idx < 5:  # 只打印前5行
                print(f"  行 {row_idx}: {row_data}")
        
        all_tables_data.append(table_data)
    
    return all_tables_data

if __name__ == '__main__':
    all_tables_data = parse_docx_content()
